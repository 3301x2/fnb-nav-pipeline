"""
apply_template.py
injects architecture content into the teams branded .docx template
preserves headers, footers, logos, brand stripe etc

usage:
    python scripts/apply_template.py docs/template.docx executive
    python scripts/apply_template.py docs/template.docx technical

requires: pip install python-docx
"""

import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# -- helpers --

def clear_body(doc):
    """Remove all paragraphs and tables from the body, keep headers/footers."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            body.remove(child)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_bold_and_normal(doc, bold_text, normal_text):
    p = doc.add_paragraph()
    b = p.add_run(bold_text)
    b.bold = True
    p.add_run(normal_text)
    return p

def add_bullet(doc, text):
    try:
        p = doc.add_paragraph(text, style='List Bullet')
    except KeyError:
        p = doc.add_paragraph()
        p.add_run(f'  \u2022 {text}')
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    pass  # table style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table

def add_callout(doc, text):
    """Add a bordered callout paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    run = p.add_run(text)
    run.italic = True
    # Add left border via XML
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    left = pBdr.makeelement(qn('w:left'), {
        qn('w:val'): 'single',
        qn('w:sz'): '24',
        qn('w:space'): '4',
        qn('w:color'): 'D97706'
    })
    pBdr.append(left)
    pPr.append(pBdr)
    return p

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break=None)
    from docx.oxml.ns import qn
    br = run._r.makeelement(qn('w:br'), {qn('w:type'): 'page'})
    run._r.append(br)


# -- executive content --

def build_executive(doc):
    clear_body(doc)

    # TL;DR
    add_heading(doc, 'TL;DR', level=1)
    add_para(doc, 'We transformed a manual, fragmented analytics environment into an automated pipeline that processes 2.2 billion transactions and 7.2 million customers to produce client-ready insights. Two machine learning models run inside BigQuery \u2014 customer segmentation and churn prediction. A prototype dashboard lets the team pitch any client by selecting from a dropdown. No recoding, no re-running.')
    add_para(doc, 'The Streamlit dashboard is a development and feedback tool only. Once stakeholders sign off on the views, the final client-facing dashboard will be built in Looker Studio as per requirements.')

    add_callout(doc, 'Note on version control: No shared or public repositories have been provisioned for this project. All source code is managed via local Git on the developer\u2019s machine using VS Code. No data, credentials, or proprietary information is stored externally. When a shared repository is provisioned, the codebase can be migrated in minutes.')

    # 1. What changed
    add_heading(doc, '1. What changed', level=1)

    add_heading(doc, 'Before', level=2)
    add_para(doc, 'For each client pitch, the team would write a query from scratch, run it in BigQuery, export results to Excel, and manually build visuals. This process repeated for every client. There were 30+ loose shared queries with no naming convention, 16+ unversioned notebooks, duplicate queries, and field names like \u201Cdemo_5\u201D that nobody could decipher without a lookup table.')

    add_heading(doc, 'After', level=2)
    add_para(doc, 'One command builds the entire pipeline. The team picks a category and client from a dropdown \u2014 Monday they pitch Adidas, Wednesday they pitch Nike, Friday they pitch Shell. Same pipeline, same dashboard, zero SQL changes. Competitors are automatically anonymized.')

    # 2. Business value
    add_heading(doc, '2. Business value delivered', level=1)

    add_heading(doc, 'Customer segmentation (ML)', level=2)
    add_para(doc, 'The pipeline automatically groups 5.8 million customers into 5 behavioral segments:')

    add_table(doc,
        ['Segment', '% customers', '% revenue', 'Avg spend', 'Action'],
        [
            ['Champions', '8.2%', '44.5%', 'R770,553', 'Retain with exclusive offers'],
            ['Loyal High Value', '14.8%', '22.3%', 'R270,103', 'Cross-sell into new categories'],
            ['Steady Mid-Tier', '30.1%', '18.7%', 'R67,662', 'Upsell to higher tiers'],
            ['At Risk', '19.9%', '8.3%', 'R11,432', 'Re-engage before they leave'],
            ['Dormant', '27.0%', '6.2%', 'R11,141', 'Win-back or accept attrition'],
        ],
        col_widths=[1.5, 0.9, 0.9, 0.9, 2.3]
    )

    add_bold_and_normal(doc, 'Headline: ', 'Champions are 8.2% of customers but drive 44.5% of revenue.')

    add_heading(doc, 'Churn prediction (ML)', level=2)
    add_para(doc, 'The model scores every customer with a churn probability (0\u2013100%):')
    add_bullet(doc, '666,301 customers (12%) flagged as Critical or High risk')
    add_bullet(doc, 'R31.6 billion in historical spend at risk')
    add_bullet(doc, 'A 10% re-engagement rate would recover ~R3.16 billion')
    add_bullet(doc, 'Model accuracy: 83.1%')

    add_heading(doc, 'Client pitch readiness', level=2)
    add_bullet(doc, 'Share of wallet: what % of customers\u2019 category spend goes to the client vs competitors')
    add_bullet(doc, 'Benchmarks: client shown by name, competitors anonymized as Competitor 1, 2, 3')
    add_bullet(doc, 'Trends: monthly client spend vs category total')
    add_bullet(doc, 'Demographics: who shops in the category by age, gender, income')
    add_bullet(doc, 'Geography: province and municipality spend concentration')

    page_break(doc)

    # 3. Dashboard strategy
    add_heading(doc, '3. Dashboard strategy', level=1)

    add_table(doc,
        ['Phase 1 (now)', 'Phase 2 (next)', 'Phase 3 (final)'],
        [
            [
                'Streamlit prototype: all 12 views built, data validated, iterate with stakeholders',
                'Stakeholder sign-off: confirm which views matter, request changes',
                'Looker Studio build: brand-compliant, creative team styles, client-facing pitch decks'
            ]
        ],
        col_widths=[2.2, 2.2, 2.2]
    )

    add_para(doc, 'The Streamlit dashboard is a development and feedback tool only. The final client-facing dashboard will be built in Looker Studio, where the creative team can own fonts, colors, and branding for pitch decks.')

    # 4. What's next
    add_heading(doc, '4. What\u2019s next', level=1)

    add_table(doc,
        ['Phase', 'What', 'Why', 'Effort'],
        [
            ['Incremental loads', 'Only process new data each run', '~90% cost reduction', '1 week'],
            ['SQL anonymization', 'Move anonymization from app into database', 'Security + Looker', '2 days'],
            ['Looker Studio', 'Build final client-facing views', 'Team standard', '2 weeks'],
            ['Dataform migration', 'Replace bash with BQ native orchestration', 'Scheduling + testing', '1 week'],
            ['POPIA compliance', 'Data classification, row-level security', 'Legal requirement', '1 week'],
        ],
        col_widths=[1.4, 2.4, 1.5, 1.2]
    )

    # 5. Scale
    add_heading(doc, '5. Scale', level=1)

    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Raw transactions processed', '2,237,540,286'],
            ['Customers profiled', '7,257,025'],
            ['Customers segmented by ML', '5,789,281'],
            ['Customers scored for churn', '5,531,237'],
            ['Destinations benchmarked', '14,536'],
        ],
        col_widths=[3.5, 3.0]
    )

    # 6. Technical summary
    add_heading(doc, '6. Technical summary', level=1)
    add_para(doc, 'The solution is a four-layer SQL pipeline (staging \u2192 features \u2192 ML \u2192 dashboard tables) running on Google BigQuery in the fmn-sandbox project, africa-south1 region. Two BigQuery ML models are trained in-warehouse: k-means clustering for customer segmentation and logistic regression for churn prediction. All tables are partitioned and clustered for cost efficiency. Infrastructure is managed by Terraform (one-command teardown). The pipeline runs via a single shell script with step-by-step execution. Four Jupyter notebooks provide interactive exploration for the analytics team.')


# -- technical content --

def build_technical(doc):
    clear_body(doc)

    # TL;DR
    add_heading(doc, 'TL;DR', level=1)
    add_para(doc, 'Full analytics pipeline on BigQuery: 2.2 billion transactions, 7.2 million customers, two ML models (k-means segmentation with Davies-Bouldin 1.29, logistic regression churn with 83% accuracy), 12-page dashboard with client dropdown and auto-anonymization. Streamlit is for dev/feedback only \u2014 final dashboard in Looker Studio. Dataform migration is ~25 hours when ready.')

    add_callout(doc, 'Note on version control and data security: No shared or public repositories have been provisioned for this project. All source code is managed via local Git on the developer\u2019s machine using VS Code. No data, credentials, PII, or proprietary information is stored externally. The repository contains only SQL scripts, Python application code, Terraform configuration, Jupyter notebooks, and documentation \u2014 zero data files. When an internal repository is provisioned, the codebase can be migrated with a single command.')

    # 1. Pipeline
    add_heading(doc, '1. Pipeline architecture', level=1)
    add_para(doc, 'Four-layer medallion architecture, executed sequentially:')

    add_table(doc,
        ['Layer', 'Tables', 'Purpose', 'Key feature'],
        [
            ['01_staging', 'stg_transactions, stg_customers', 'Clean, join lookups, strip PII, rename fields', 'Partitioned by month, clustered by category + destination'],
            ['02_intermediate', 'int_rfm_features, int_rfm_scores, int_customer_category_spend, int_destination_metrics', 'RFM features, quintile scores, spend share, destination KPIs', 'Share-of-wallet per customer \u00D7 category \u00D7 destination'],
            ['03_ml', 'kmeans_customer_segments (MODEL), churn_classifier (MODEL), mart_cluster_output, mart_churn_risk', 'K-means (k=5), logistic regression churn', 'BigQuery ML: in-warehouse training and inference'],
            ['04_marts', '8 dashboard-ready tables', 'Profiles, summaries, trends, benchmarks, geo, demographics', 'Dashboard reads from marts only'],
        ],
        col_widths=[1.2, 1.8, 1.8, 1.7]
    )

    # 2. Execution
    add_heading(doc, '2. Execution model', level=1)
    add_para(doc, 'Step-by-step with environment selection:')
    add_bullet(doc, 'bash scripts/run.sh sandbox \u2014 all steps on fmn-sandbox')
    add_bullet(doc, 'bash scripts/run.sh production 3 \u2014 step 3 on fmn-production')
    add_bullet(doc, 'bash scripts/run.sh sandbox 1 \u2014 step 1 on fmn-sandbox')
    add_para(doc, 'SQL files use a __PROJECT__ placeholder. The runner substitutes the actual project ID at runtime \u2014 same code targets sandbox or production.')

    # 3. ML
    add_heading(doc, '3. ML models', level=1)

    add_table(doc,
        ['Model', 'Type', 'Features', 'Key metric', 'Output'],
        [
            ['Customer segmentation', 'K-means (unsupervised)', '9 RFM features', 'Davies-Bouldin: 1.295', '5 segments: Champions \u2192 Dormant'],
            ['Churn prediction', 'Logistic regression (supervised)', '15 features', 'Accuracy: 0.831, F1: 0.556', 'Probability per customer (0\u2013100%)'],
        ],
        col_widths=[1.3, 1.3, 1.1, 1.3, 1.5]
    )

    add_para(doc, 'Both models run inside BigQuery ML (in-warehouse). No data egress, no external compute. The churn model uses logistic regression because boosted tree classifiers are not yet supported in the africa-south1 region.')

    # 4. Dashboard
    add_heading(doc, '4. Dashboard strategy', level=1)
    add_bold_and_normal(doc, 'Streamlit is the development and feedback tool. ', 'Looker Studio is the final deliverable.')

    add_table(doc,
        ['Phase 1: Streamlit (now)', 'Phase 2: Sign-off', 'Phase 3: Looker Studio'],
        [['12 draft pages, git-controlled, iterate in real-time', 'Stakeholders confirm views, metrics, and filters', 'Brand-compliant, creative team owns styling, client-facing']],
        col_widths=[2.2, 2.2, 2.2]
    )

    add_para(doc, 'Dashboard targets the project via environment variable: BQ_PROJECT=fmn-production streamlit run dashboards/app.py')

    page_break(doc)

    # 5. Dataform
    add_heading(doc, '5. Dataform comparison', level=1)
    add_para(doc, 'Dataform is BigQuery\u2019s native transformation framework. Migration is mechanical \u2014 every SELECT statement stays the same.')

    add_table(doc,
        ['Aspect', 'Current (SQL + bash)', 'Dataform'],
        [
            ['Table creation', 'CREATE OR REPLACE in every .sql file', 'Dataform generates DDL from config block'],
            ['Dependencies', 'Script execution order (steps 0\u20135)', 'Automatic DAG from ref() calls'],
            ['Project ID', '__PROJECT__ placeholder, replaced at runtime', 'Set once in dataform.json'],
            ['Incremental', 'Custom DECLARE/IF/INSERT per file', 'Built-in is_incremental() macro'],
            ['Testing', 'Separate validate.sh script', 'Native assertions: uniqueKey, nonNull'],
            ['Scheduling', 'Manual: bash scripts/run.sh', 'Built-in cron in GCP console'],
            ['Documentation', 'Separate data_dictionary.md', 'Inline descriptions, auto-generated catalog'],
        ],
        col_widths=[1.4, 2.5, 2.6]
    )

    # 6. Risk
    add_heading(doc, '6. Risk matrix', level=1)

    add_table(doc,
        ['Risk', 'SQL + bash', 'Dataform', 'Current mitigation'],
        [
            ['Partial failure', 'High: re-run scans all data', 'Low: DAG retries failed nodes', 'Manual re-run from failed step'],
            ['Cost overrun', 'High: 2.2B rows scanned', 'Low: incremental only', 'Incremental logic (next phase)'],
            ['Wrong order', 'Medium: script enforces', 'None: ref() enforces DAG', 'Step-by-step runner'],
            ['Data quality', 'Medium: post-hoc checks', 'Low: assertions block downstream', '5 checks in validate.sh'],
            ['Onboarding', 'Medium: read run.sh', 'Low: visual DAG', 'README + data dictionary'],
        ],
        col_widths=[1.2, 1.8, 1.8, 1.7]
    )

    # 7. Migration
    add_heading(doc, '7. Migration path (~25 hours)', level=1)

    add_table(doc,
        ['Task', 'Files', 'Effort', 'Complexity'],
        [
            ['Initialise Dataform repo', '1', '1 hour', 'Low'],
            ['Convert staging to .sqlx', '2', '2 hours', 'Low'],
            ['Convert intermediate to .sqlx', '4', '3 hours', 'Low'],
            ['Convert ML to operations', '4', '4 hours', 'Medium'],
            ['Convert marts to .sqlx', '8', '3 hours', 'Low'],
            ['Add assertions', '16', '4 hours', 'Low'],
            ['Add incremental logic', '3', '4 hours', 'Medium'],
            ['Testing + handoff', 'All', '4 hours', 'Low'],
            ['Total', '16 files', '~25 hours', 'Low\u2013Medium'],
        ],
        col_widths=[2.4, 1.0, 1.0, 2.1]
    )

    page_break(doc)

    # 8. Production results
    add_heading(doc, '8. Production results', level=1)

    add_heading(doc, 'Scale', level=2)
    add_table(doc,
        ['Metric', 'Value'],
        [
            ['Raw transactions processed', '2,237,540,286'],
            ['Customers staged', '7,257,025'],
            ['Customers segmented (k-means)', '5,789,281'],
            ['Customers scored (churn ML)', '5,531,237'],
            ['Customer-category-destination rows', '212,809,215'],
            ['Destinations benchmarked', '14,536'],
        ],
        col_widths=[3.5, 3.0]
    )

    add_heading(doc, 'ML model performance', level=2)
    add_table(doc,
        ['Metric', 'K-means segmentation', 'Churn classifier'],
        [
            ['Model type', 'K-means (unsupervised)', 'Logistic regression (supervised)'],
            ['Davies-Bouldin', '1.295 (< 2.0 = good)', 'N/A'],
            ['Accuracy', 'N/A', '0.831'],
            ['F1 score', 'N/A', '0.556'],
            ['Cluster balance', 'Largest: 30.1% (< 40%)', 'N/A'],
            ['Business insight', '8.2% of customers drive 44.5% of revenue', '666K at risk = R31.6B spend'],
        ],
        col_widths=[1.8, 2.4, 2.3]
    )

    # 9. Recommendation
    add_heading(doc, '9. Recommendation', level=1)
    add_para(doc, 'The current SQL + bash implementation is appropriate for the current phase: initial deployment, validation, and stakeholder buy-in.')
    add_bold_and_normal(doc, 'Dataform migration is recommended as the next phase, ', 'not because the current approach is flawed, but because the pipeline has proven its value and now needs production hardening.')
    add_para(doc, 'Trigger migration when:')
    add_bullet(doc, 'The pipeline needs a recurring schedule (daily/weekly)')
    add_bullet(doc, 'A second developer joins the project')
    add_bullet(doc, 'BigQuery scan costs exceed threshold due to full refreshes')
    add_bullet(doc, 'Dataform access is provisioned for the team')

    # 10. File inventory
    add_heading(doc, '10. File inventory', level=1)
    add_table(doc,
        ['File', 'Purpose'],
        [
            ['sql/01_staging/ (2 files)', 'Join lookups, partition, cluster, strip PII, rename demo_*'],
            ['sql/02_intermediate/ (4 files)', 'RFM features, quintile scores, spend share, destination KPIs'],
            ['sql/03_ml/ (4 files)', 'K-means training + predict, churn training + predict'],
            ['sql/04_marts/ (8 files)', 'Dashboard-ready analytical tables'],
            ['dashboards/app.py', '12-page Streamlit dashboard (dev/feedback only)'],
            ['notebooks/ (4 files)', 'Exploration, features, cluster profiling, client pitch'],
            ['terraform/main.tf', 'Infrastructure as code (create/destroy)'],
            ['scripts/', 'Pipeline runner, validation, Cloud Run deployment'],
            ['docs/', 'Architecture docs, data dictionary, Word template'],
        ],
        col_widths=[2.8, 3.7]
    )


# -- main --

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python scripts/apply_template.py docs/template.docx executive")
        print("  python scripts/apply_template.py docs/template.docx technical")
        print("  python scripts/apply_template.py docs/template.docx both")
        sys.exit(1)

    template_path = sys.argv[1]
    doc_type = sys.argv[2].lower()

    if doc_type in ('executive', 'exec', 'both'):
        print(f"Building executive doc from {template_path}...")
        doc = Document(template_path)
        build_executive(doc)
        out = 'docs/architecture_executive_branded.docx'
        doc.save(out)
        print(f"  \u2713 Saved: {out}")

    if doc_type in ('technical', 'tech', 'both'):
        print(f"Building technical doc from {template_path}...")
        doc = Document(template_path)
        build_technical(doc)
        out = 'docs/architecture_technical_branded.docx'
        doc.save(out)
        print(f"  \u2713 Saved: {out}")

    print("\nDone. The output inherits all template styles:")
    print("  - Fonts, heading styles, spacing")
    print("  - Headers and footers (logos, brand stripe)")
    print("  - Page margins and layout")
